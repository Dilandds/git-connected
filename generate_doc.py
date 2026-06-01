from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── helpers ──────────────────────────────────────────────────────────────────

def add_main_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)

def add_screen_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)

def add_section(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x37, 0x47, 0x51)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)

def add_row(doc, question, answer):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.2)
    q_run = p.add_run(f"{question}  ")
    q_run.bold = True
    q_run.font.size = Pt(10)
    a_run = p.add_run(answer)
    a_run.font.size = Pt(10)
    a_run.font.color.rgb = RGBColor(0x16, 0x7A, 0x3B)

def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"⚠  {text}")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.2)

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 80)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# ── HEADER ───────────────────────────────────────────────────────────────────

add_main_title(doc, "Project Requirements Summary")
p = doc.add_paragraph("Based on buyer answers — 01 June 2026")
p.runs[0].font.size = Pt(10)
p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
add_divider(doc)

# ── USER ROLES ───────────────────────────────────────────────────────────────

add_screen_title(doc, "User Roles")
add_row(doc, "Total users:", "Maximum 5 — 1 Admin (creator) + up to 4 collaborators")
add_row(doc, "Who sets passwords:", "Admin creator sets passwords for all other users")
add_row(doc, "Admin powers:", "Full access — adds/removes users, edits all content, protects sections")
add_row(doc, "Collaborator access:", "Limited — can view and interact but not edit core data")
add_note(doc, "No role names defined yet — clarify if collaborators all have the same permissions or differ per user")
add_divider(doc)

# ── TIMELINE ─────────────────────────────────────────────────────────────────

add_screen_title(doc, "Timeline")

add_section(doc, "Project Info")
add_row(doc, "Fields editable by:", "Admin creator only")
add_row(doc, "+Add photo:", "Yes — project cover image")

add_section(doc, "Operators")
add_row(doc, "What is an Operator:", "External collaborator / supplier company")
add_row(doc, "Operator name:", "Can be renamed to the company name")
add_row(doc, "Who adds operators:", "Admin creator only — admin enters all operator info")
add_row(doc, "What operators see:", "Their own lane only")
add_note(doc, "Operators are external — they log in and see only their tasks, not the full project timeline")

add_section(doc, "Operations & Tasks")
add_row(doc, "Sub-tasks:", "Yes — e.g. Operation = 'Setting diamonds', sub-tasks = individual components")
add_row(doc, "Who adds/edits operations:", "Admin creator only")
add_row(doc, "Urgent flag:", "Can be set on any task")
add_row(doc, "Task block movement:", "Snap to days — set by defining start/end dates")
add_row(doc, "Multiple operators per task:", "No — each operator has their own operation")
add_row(doc, "Editing a task:", "Click a date block → form opens → fill in → close form")

add_section(doc, "Status & Legend")
add_row(doc, "Who changes status:", "Admin creator only")
add_row(doc, "Custom legend types:", "Shared across the whole project")

add_section(doc, "Comments")
add_row(doc, "Comment type:", "Full thread of comments per task")
add_row(doc, "Visible to:", "All users")

add_section(doc, "Bottom Panel")
add_row(doc, "Duration:", "Auto-calculated from start and end dates")
add_row(doc, "Deadline:", "A red vertical line placed directly on the timeline — date shown automatically. Each task also has its own deadline via the task form")
add_row(doc, "Contributors:", "Per task — selected from saved user list")
add_row(doc, "Project Manager / Technical Manager:", "Selected from saved user list")
add_row(doc, "Clicking a task:", "Yes — always opens details in bottom panel")

add_section(doc, "Overview Tab")
add_row(doc, "Interactable:", "Read-only — auto-mirrors all operator tabs")
add_row(doc, "Week / Month views:", "Block colours only — no task names (too small)")

add_section(doc, "Export")
add_row(doc, "PDF export:", "Yes")
add_divider(doc)

# ── VALIDATION ───────────────────────────────────────────────────────────────

add_screen_title(doc, "Validation")

add_section(doc, "General")
add_row(doc, "Who creates validation session:", "All project managers")
add_row(doc, "Session date:", "Automatic with option to modify manually")
add_row(doc, "Editing past sessions:", "Locked after approval")
add_row(doc, "Approval mechanism:", "Internal digital signature — Approved / Rejected / Needs correction")
add_row(doc, "Links to Version Comparison:", "No")
add_note(doc, "Digital signature implies a formal sign-off flow — clarify who can sign (admin only or all project managers)")

add_section(doc, "Stakeholders Table")
add_row(doc, "Who adds/removes contributors:", "Admin creator")
add_row(doc, "Role names:", "Customisable — saved in drop-down list")
add_row(doc, "Name/Company field:", "Customisable — saved in drop-down list")
add_row(doc, "Who updates Status & Progress:", "Admin creator")
add_row(doc, "Deliverables:", "Both file attachments and text descriptions")
add_row(doc, "Progress %:", "Auto-calculated")

add_section(doc, "Presentation Elements")
add_row(doc, "Photos:", "Uploaded fresh each time")
add_row(doc, "Photo label:", "Customisable")

add_section(doc, "Estimated Cost & Schedule in Validation")
add_row(doc, "Cost data:", "Auto-pulled from Best partner in Estimated Cost screen")
add_row(doc, "Lead time:", "Auto-pulled from Timeline")
add_row(doc, "Schedule checkboxes:", "Auto-updated from Timeline")
add_row(doc, "Who can edit:", "Admin only")

add_section(doc, "Validation Report")
add_row(doc, "Who fills in:", "Admin creator")
add_row(doc, "Decision Maker / Presentation Lead:", "Customisable drop-down list (saved)")
add_row(doc, "Modifications table — Responsible:", "Manual — Priority/Status drop list: To modify / Accepted / On hold / Canceled")
add_row(doc, "Action Plan — Responsible:", "Customisable drop-down list (saved)")
add_row(doc, "Add/remove rows:", "Yes")
add_row(doc, "Locked after submit:", "Yes")

add_section(doc, "Notifications")
add_row(doc, "Stakeholder added:", "No notification")
add_row(doc, "Action assigned to user:", "Yes — notified")

add_section(doc, "Export")
add_row(doc, "PDF export:", "Yes")
add_divider(doc)

# ── REPORT ───────────────────────────────────────────────────────────────────

add_screen_title(doc, "Report")

add_section(doc, "General")
add_row(doc, "Who creates report:", "Admin creator")
add_row(doc, "Report date:", "Manually selected via calendar")
add_row(doc, "Editing past reports:", "Locked permanently after save")

add_section(doc, "Header")
add_row(doc, "Company logo:", "Uploaded once — reused across all reports")
add_row(doc, "Launch deadline:", "Entered manually")
add_row(doc, "Project name / Reference:", "Auto-filled from project info")

add_section(doc, "Company & Partners")
add_row(doc, "Project Manager / Technical Manager / Quality Lead:", "Selected from customisable saved list")
add_row(doc, "Partners:", "Entered manually")
add_row(doc, "Other rows:", "Customisable")

add_section(doc, "Present Attendees")
add_row(doc, "Column headers:", "Customisable")
add_row(doc, "Attendee names:", "Selected from customisable saved drop-down list")

add_section(doc, "Content")
add_row(doc, "Production follow-up text:", "Plain text")
add_row(doc, "Photo rows:", "Can be increased (unlimited)")
add_row(doc, "Photos:", "Uploaded fresh each time")
add_row(doc, "Caption below photo:", "Yes — customisable text field")

add_section(doc, "Pages")
add_row(doc, "Add pages:", "Yes — unlimited")
add_row(doc, "Section layout:", "Can be rearranged")

add_section(doc, "Export")
add_row(doc, "Print/PDF:", "Both — export as PDF or send to printer")
add_row(doc, "Merge reports:", "Yes — can merge multiple reports or export individually")
add_divider(doc)

# ── ESTIMATED COST ───────────────────────────────────────────────────────────

add_screen_title(doc, "Estimated Cost")

add_section(doc, "General")
add_row(doc, "What is a Trade:", "The activity/scope of work of a partner (e.g. manufacturing, design)")
add_row(doc, "Max trades:", "Unlimited — add with + button")
add_row(doc, "Who adds/removes trades:", "Admin creator")

add_section(doc, "Partners")
add_row(doc, "Max partners per trade:", "10")
add_row(doc, "Partner names:", "Customisable drop-down list (saved)")
add_row(doc, "Who adds/removes partners:", "Admin creator")
add_row(doc, "Best partner selection:", "Manually chosen by admin creator")
add_row(doc, "Edit after Best is set:", "Yes — can still be modified")

add_section(doc, "Task Details")
add_row(doc, "Component / Task fields:", "Manually entered")
add_row(doc, "Add/remove rows:", "Yes")
add_row(doc, "Currency:", "Global per project — USD, EUR, or any world currency")
add_row(doc, "Hourly rate:", "Entered once in the first row — auto-filled down the entire column")
add_note(doc, "Hourly rate auto-fill down column is a key behaviour to build carefully")

add_section(doc, "Comparison Tab")
add_row(doc, "Populated:", "Auto-populated from Partner tabs")
add_row(doc, "Editable:", "No — read-only")

add_section(doc, "Overview Tab")
add_row(doc, "Populated:", "Auto-generated from all trades and their Best partners")
add_row(doc, "Access:", "Password-protected by admin creator")
add_note(doc, "Estimated Cost section has a separate password set by the admin — clarify if this is per-project or global")

add_section(doc, "Export")
add_row(doc, "Print/PDF:", "Both — export as PDF or send to printer")
add_row(doc, "Per trade export:", "Yes — individual trades can be exported separately")
add_divider(doc)

# ── FILES AND VERSIONS ───────────────────────────────────────────────────────

add_screen_title(doc, "Files and Versions")
add_note(doc, "No answers received yet for this screen — questions still pending")
add_divider(doc)

# ── VERSION COMPARISON ───────────────────────────────────────────────────────

add_screen_title(doc, "Version Comparison")
add_note(doc, "No answers received yet for this screen — questions still pending")
add_divider(doc)

# ── PROJECT BRIEF ────────────────────────────────────────────────────────────

add_screen_title(doc, "Project Brief")
add_note(doc, "No answers received yet for this screen — questions still pending")
add_divider(doc)

# ── TRACEABILITY ─────────────────────────────────────────────────────────────

add_screen_title(doc, "Traceability")
add_note(doc, "No answers received yet for this screen — questions still pending")
add_divider(doc)

# ── KEY OBSERVATIONS ─────────────────────────────────────────────────────────

add_screen_title(doc, "Key Observations & Flags")

observations = [
    ("Admin-heavy system",
     "Almost all create/edit actions are restricted to the admin creator. Collaborators/operators have very limited write access."),
    ("Max 5 users",
     "1 admin + 4 collaborators. This is a small team tool, not an enterprise platform."),
    ("Operators = external suppliers",
     "Operators are outside companies. They log in and see only their own lane. Admin enters all their data."),
    ("Saved drop-down lists everywhere",
     "Many fields use customisable saved lists (roles, names, partners, attendees). A global list management interface will be needed."),
    ("Password-protected Estimated Cost",
     "The entire Estimated Cost section has a separate password. Clarify: is this one global PW or per project?"),
    ("Digital signature on Validation",
     "Validation uses internal sign-off (Approved / Rejected / Needs correction). Clarify who can sign and whether all signatures are needed before locking."),
    ("Hourly rate column auto-fill",
     "In Estimated Cost, entering the hourly rate once fills all rows in that column automatically. Key behaviour to implement correctly."),
    ("Deadline = red vertical line on Timeline",
     "Deadline is a draggable red line on the timeline grid, not just a date field. Each task also has its own deadline inside its form."),
    ("Files & Versions, Version Comparison, Project Brief, Traceability",
     "No answers received yet for these 4 screens. Follow up needed before scoping those modules."),
]

for title, detail in observations:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.2)
    t = p.add_run(f"{title}:  ")
    t.bold = True
    t.font.size = Pt(10)
    d = p.add_run(detail)
    d.font.size = Pt(10)

# ── SAVE ─────────────────────────────────────────────────────────────────────

output_path = "/Users/dilanweerasinghe/Desktop/Project_Requirements_Summary.docx"
doc.save(output_path)
print(f"Saved → {output_path}")
